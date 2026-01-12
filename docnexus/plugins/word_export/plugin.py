import logging
import io
import shutil
from pathlib import Path

# Note: Feature, FeatureType, FeatureState, PluginRegistry are INJECTED by the loader.
# Do not import them directly to avoid split-brain issues.

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

logger = logging.getLogger(__name__)

# Constants
MAX_EXPORT_HTML_SIZE = 50 * 1024 * 1024  # 50 MB

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph in a Word document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    
    # Insert bookmark
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)

def export_to_word(html_content: str) -> bytes:
    """
    Exports HTML content to a Word (.docx) file byte stream.
    """
    try:
        from htmldocx import HtmlToDocx
        from docx import Document
        from docx.shared import RGBColor, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        logger.error(f"Failed to import Word export dependencies: {e}")
        raise RuntimeError("Word export dependencies (htmldocx, python-docx) not installed.")

    # Size Check
    html_size = len(html_content.encode('utf-8'))
    if html_size > MAX_EXPORT_HTML_SIZE:
        raise ValueError(f"Content too large ({html_size/1024/1024:.2f} MB). Max {MAX_EXPORT_HTML_SIZE/1024/1024} MB.")

    logger.info(f"WordExport: Generating document from {html_size} bytes of HTML...")

    # Pre-process HTML with BeautifulSoup
    try:
        soup = BeautifulSoup(html_content, 'lxml')
    except:
        soup = BeautifulSoup(html_content, 'html.parser')
    
    # Cleaning (Scripts, Styles, Nav)
    for tag in soup.find_all(['script', 'style', 'nav']):
        tag.decompose()
    
    # Main Content Extraction
    # We want to include the Table of Contents (.toc-container) AND the Markdown Content (.markdown-content)
    # The frontend wraps both in #documentContent div (Line 729 view.html)
    # But usually sending full <html>.
    
    container = soup.find(id='documentContent')
    selected_content = []
    
    if container:
        # Extract TOC if present
        toc = container.find(class_='toc-container')
        if toc:
             # Style TOC for Word
            toc_header = toc.find(class_='toc-header')
            if toc_header:
                toc_header.name = 'h2' # Make it a standard header for Word
                toc_header['style'] = 'font-size: 14pt; color: #4b5563; margin-top: 0;'
            
            selected_content.append(toc)
            
        # Extract Markdown Content
        md_content = container.find(class_='markdown-content')
        if md_content:
             selected_content.append(md_content)
    else:
        # Fallback to old behavior if ID not found
        md_content = soup.find(class_='markdown-content')
        if md_content:
             selected_content.append(md_content)

    if selected_content:
        # Style Tables for Word (Apply to all tables in selected content)
        for part in selected_content:
            for table in part.find_all('table'):
                table['style'] = 'border-collapse: collapse; width: 100%; border: 2px solid rgba(99, 102, 241, 0.2); margin-bottom: 20px;'
                table['border'] = '1'
                
                # Thead check
                thead = table.find('thead')
                if not thead:
                    first_row = table.find('tr')
                    if first_row and first_row.find('th'):
                        thead = soup.new_tag('thead')
                        first_row.extract()
                        thead.append(first_row)
                        table.insert(0, thead)
                
                # Colors and Styles injection
                for th in table.find_all('th'):
                    th['bgcolor'] = '#6366f1'
                    th['style'] = 'background-color: #6366f1 !important; color: #ffffff !important;'
                
                for td in table.find_all('td'):
                    td['style'] = 'padding: 8px; border: 1px solid #e5e7eb;'
        
        # Combine content
        combined_html = "".join([str(tag) for tag in selected_content])
        clean_html = f'<html><head><meta charset="utf-8"></head><body>{combined_html}</body></html>'
        
        # Capture main_content for booking logic later (use md_content reference)
        main_content = md_content if 'md_content' in locals() and md_content else None
    else:
        # Absolute fallback
        clean_html = f'<html><body>{soup.body.decode_contents() if soup.body else str(soup)}</body></html>'
        main_content = None

    # Pre-process HTML to resolve/fetch images (crucial for stability)
    # This prevents htmldocx from crashing on network errors or missing files.
    soup = BeautifulSoup(clean_html, 'html.parser')
    
    # We need a temp dir for downloaded images that persists during conversion
    import tempfile
    import urllib.request
    from urllib.parse import urlparse
    import shutil
    
    # Create a temporary directory for this export session
    with tempfile.TemporaryDirectory() as temp_img_dir:
        temp_dir_path = Path(temp_img_dir)
        
        for img in soup.find_all('img'):
            src = img.get('src')
            if not src:
                continue
                
            is_external = src.startswith(('http://', 'https://'))
            new_src = None
            
            try:
                if is_external:
                    # Check for SVG in URL before downloading to save time
                    path_obj = Path(urlparse(src).path)
                    if path_obj.suffix.lower() == '.svg':
                         raise ValueError("SVG format is not supported by Word.")

                    # Attempt to download external image using standard library
                    req = urllib.request.Request(src, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=3.0) as response:
                        # Check Content-Type for SVG
                        ctype = response.info().get_content_type().lower()
                        if 'svg' in ctype:
                             raise ValueError("SVG format is not supported by Word.")

                        # Determine filename
                        filename = path_obj.name or "image.png"
                        local_path = temp_dir_path / filename
                        
                        with open(local_path, 'wb') as f:
                            shutil.copyfileobj(response, f)
                    
                    new_src = str(local_path)
                    
                elif src.startswith('data:image/'):
                    # Handle Data URIs (e.g. Mermaid Exports)
                    import base64
                    try:
                        from PIL import Image
                    except ImportError:
                        Image = None
                    
                    if ';base64,' in src:
                        header, data = src.split(';base64,')
                        ctype = header.split(':')[1]
                        
                        if 'svg' in ctype:
                             raise ValueError("SVG data URIs are not supported.")
                        
                        img_data = base64.b64decode(data)
                        
                        # Process Image (Flatten Transparency)
                        if Image:
                            try:
                                with Image.open(io.BytesIO(img_data)) as im:
                                   # Convert to RGBA if strictly RGB to ensure consistent handling (though usually PNG is RGBA)
                                   if im.mode in ('RGBA', 'LA') or (im.mode == 'P' and 'transparency' in im.info):
                                       # Create white background
                                       bg = Image.new('RGB', im.size, (255, 255, 255))
                                       # Paste image on top using alpha channel
                                       if im.mode != 'RGBA':
                                           im = im.convert('RGBA')
                                       bg.paste(im, mask=im.split()[3]) # 3 is the alpha channel
                                       
                                       # Save flattened image
                                       output = io.BytesIO()
                                       bg.save(output, format='PNG')
                                       img_data = output.getvalue()
                                       ext = '.png'
                                   else:
                                        ext = '.png' if 'png' in ctype else '.jpg'
                            except Exception as iconv_err:
                                logger.warning(f"PIL Conversion failed, using original data: {iconv_err}")
                                ext = '.png' if 'png' in ctype else '.jpg'
                        else:
                            ext = '.png' if 'png' in ctype else '.jpg'

                        filename = f"embedded_image_{hash(data)}{ext}"
                        local_path = temp_dir_path / filename
                        
                        with open(local_path, 'wb') as f:
                            f.write(img_data)
                            
                        new_src = str(local_path)
                    else:
                        raise ValueError("Unsupported Data URI format")

                elif not src.startswith('data:'):
                     # Resolve local path
                     # Try relative to CWD first
                    potential_path = Path(src).resolve()
                    if not potential_path.exists():
                        # Try relative to server root if CWD failed
                        potential_path = (Path(os.getcwd()) / src.lstrip('/\\')).resolve()
                    
                    if potential_path.exists():
                        if potential_path.suffix.lower() == '.svg':
                             raise ValueError("SVG format is not supported.")
                        new_src = str(potential_path)
                    else:
                        logger.warning(f"Word Export: Local image not found: {src}")

                # Update src if valid path found
                if new_src:
                    img['src'] = new_src
                else:
                    raise ValueError("Could not resolve image source.")

            except Exception as e:
                # logger.warning(f"Word Export: Skipping image '{src}': {e}")
                # Fallback to alt text
                alt_text = img.get('alt', 'Image')
                replacement = soup.new_tag('span')
                replacement.string = f"[{alt_text}]"
                replacement['style'] = "color: #666; font-style: italic; border: 1px solid #ccc; padding: 2px;"
                img.replace_with(replacement)

        # Sanitize Styles to prevent htmldocx crashes (invalid literal for int() with base 16)
        # htmldocx chokes on 'stroke: none', 'fill: ...', and 'color: auto/none' often found in Mermaid/Shims
        for tag in soup.find_all(True):
            if tag.has_attr('style'):
                styles = [s.strip() for s in tag['style'].split(';') if s.strip()]
                clean_styles = []
                for s in styles:
                    if ':' in s:
                        prop, val = s.split(':', 1)
                        prop = prop.strip().lower()
                        val = val.strip().lower()
                        
                        # Strip dangerous SVG-related styles that htmldocx doesn't understand
                        if prop in ['stroke', 'stroke-width', 'fill', 'fill-opacity', 'stroke-opacity']:
                            continue
                            
                        # Strip color values that aren't strict hex/rgb (causes base 16 error)
                        if 'color' in prop: # color, background-color, border-color
                             if val in ['none', 'auto', 'transparent', 'inherit', 'initial', 'unset']:
                                 continue
                             # Try to protect against 'rgba(0,0,0,0)' if htmldocx doesn't support it (it usually doesn't)
                             if 'rgba' in val:
                                 continue

                        clean_styles.append(s)
                
                if clean_styles:
                    tag['style'] = "; ".join(clean_styles)
                else:
                    del tag['style']
                    
        # Also remove 'stroke' and 'fill' attributes directly
        for tag in soup.find_all(attrs={"stroke": True}):
            del tag['stroke']
        for tag in soup.find_all(attrs={"fill": True}):
            del tag['fill']
        for tag in soup.find_all(attrs={"viewbox": True}):
             del tag['viewbox'] # Clean up any lingering SVG debris

        clean_html = str(soup)

        # Generate Word Doc
        doc = Document()
        new_parser = HtmlToDocx()
        
        try:
            # Now safe to convert
            new_parser.add_html_to_document(clean_html, doc)
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            logger.error(f"HtmlToDocx conversion failed: {e}", exc_info=True)
            doc.add_paragraph(f"[Export Error: Document content could not be fully converted.]")
            doc.add_paragraph(f"Details: {str(e)}")
            # Add traceback in small font for debugging
            p = doc.add_paragraph()
            run = p.add_run(error_details)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Post-processing (Bookmarks)
        heading_ids = {}
        if main_content:
            for heading in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                 if heading.get('id'):
                     heading_ids[heading.get_text(strip=True)] = heading.get('id')
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and paragraph.text.strip() in heading_ids:
                add_bookmark(paragraph, heading_ids[paragraph.text.strip()])

        # Post-processing (Image Sizing & Centering)
        # Fixes oversized diagrams in Word export
        try:
            # Calculate writable limits
            section = doc.sections[0]
            page_width = section.page_width
            page_height = section.page_height
            margin_x = section.left_margin + section.right_margin
            margin_y = section.top_margin + section.bottom_margin
            
            writable_width = page_width - margin_x
            writable_height = page_height - margin_y
            
            from docx.shared import Emu

            for shape in doc.inline_shapes:
                # Calculate aspect ratio
                if shape.width == 0: continue
                aspect_ratio = shape.height / shape.width
                
                # 1. Width Constraint
                if shape.width > writable_width:
                    shape.width = writable_width
                    shape.height = int(writable_width * aspect_ratio)
                
                # 2. Height Constraint (Applied after width to ensure final fit)
                if shape.height > writable_height:
                    shape.height = writable_height
                    shape.width = int(writable_height / aspect_ratio)
                
                # Force Center Alignment for paragraphs containing images
                # This works because htmldocx usually puts block images in their own p (or we forced it)
                # We check if the paragraph is mostly just this image to avoid centering mixed content
                # For now, simplistic approach: if paragraph has an inline shape, center it.
                # Accessing the parent paragraph of a shape isn't direct in python-docx public API,
                # but we can try iterating paragraphs and finding runs with drawings.
                pass 
            
            # Robust Centering Loop
            # Iterate paragraphs, find those with images, force center
            for p in doc.paragraphs:
                # Check for blip/drawing
                if 'Graphic' in p._element.xml or 'drawing' in p._element.xml:
                    # Heuristic: mostly image content?
                    if len(p.text.strip()) < 5:  # Almost no text, valid assumption for a diagram block
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
        except Exception as e:
            logger.warning(f"Failed to resize images: {e}")

        # Post-processing (Style Table Grid)
        for table in doc.tables:
            table.style = 'Table Grid'
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '6366f1')
                    cell._element.get_or_add_tcPr().append(shading_elm)

        # Post-processing (Fix Internal Hyperlinks for TOC)
        # htmldocx creates external links for #anchors. We need to convert them to w:anchor.
        try:
            part = doc.part
            rels = part.rels
            
            # Iterate all paragraphs to find hyperlinks
            for p in doc.paragraphs:
                for child in p._element:
                    if child.tag.endswith('hyperlink'):
                        rid = child.get(qn('r:id'))
                        if rid and rid in rels:
                            rel = rels[rid]
                            if rel.target_ref and rel.target_ref.startswith('#'):
                                # Found internal link candidate
                                anchor_name = rel.target_ref[1:] # strip #
                                
                                # Convert to internal anchor
                                # Note: We assume the bookmark name matches the ID (which we ensured in our bookmarking logic? 
                                # Actually our bookmark logic uses get_text()??
                                # NO. 
                                # Line 233: heading_ids[heading.get_text()] = heading.get('id')
                                # Line 237: add_bookmark(..., heading_ids[...]) -> creates bookmark with NAME = ID.
                                # So if href="#id", and bookmark name="id", it works.
                                
                                child.set(qn('w:anchor'), anchor_name)
                                # Remove r:id (external reference)
                                try:
                                    del child.attrib[qn('r:id')]
                                except:
                                    pass
                                child.set(qn('w:history'), '1')
        except Exception as e:
            logger.warning(f"Failed to fix internal hyperlinks: {e}")

        # Save to Buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
    # Context exits, temp dir deleted.
    # Context exits, temp dir deleted.
    logger.info("WordExport: Complete. Returning bytes.")
    return buffer.getvalue()

# Expose features for FeatureManager
def get_features():
    # Helper to access injected classes safely
    # If not injected (e.g. static analysis), these might fail, which is expected.
    _Feature = globals().get('Feature')
    _FeatureType = globals().get('FeatureType')
    _FeatureState = globals().get('FeatureState')
    
    if not _Feature:
        logger.error("Plugin dependency injection failed: Feature class missing.")
        return []

    return [
        _Feature(
            name="docx",
            handler=export_to_word,
            feature_type=_FeatureType.EXPORT_HANDLER,
            state=_FeatureState.STANDARD
        )
    ]

# Metadata
PLUGIN_METADATA = {
    'name': 'Word Export',
    'description': 'Exports documentation to Microsoft Word (.docx) with TOC and styles.',
    'category': 'export',
    'icon': 'fa-file-word',
    'preinstalled': True
}
